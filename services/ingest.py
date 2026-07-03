"""Extract plain text from PDF, EPUB, and text uploads."""

import html
import re
import zipfile
from pathlib import Path


def _html_to_text(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", "", raw)
    raw = re.sub(r"(?i)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</p\s*>", "\n\n", raw)
    raw = re.sub(r"(?i)</h[1-6]\s*>", "\n\n", raw)
    raw = re.sub(r"<[^>]+>", "", raw)
    return html.unescape(raw)


def epub_to_text(epub_path: Path) -> str:
    try:
        import ebooklib
        from ebooklib import epub
    except ImportError as exc:
        raise RuntimeError("ebooklib required for EPUB support") from exc

    book = epub.read_epub(str(epub_path))
    parts: list[str] = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        content = item.get_content().decode("utf-8", errors="ignore")
        text = _html_to_text(content)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _epub_fallback_zip(epub_path: Path) -> str:
    """Minimal EPUB parser without ebooklib."""
    parts: list[str] = []
    with zipfile.ZipFile(epub_path) as zf:
        html_names = sorted(
            n for n in zf.namelist()
            if n.lower().endswith((".xhtml", ".html", ".htm")) and "nav" not in n.lower()
        )
        for name in html_names:
            raw = zf.read(name).decode("utf-8", errors="ignore")
            text = _html_to_text(raw).strip()
            if text:
                parts.append(text)
    if not parts:
        raise ValueError("No readable content in EPUB")
    return "\n\n".join(parts)


def pdf_to_text(pdf_path: Path) -> str:
    """Extract embedded text from a digital PDF (no OCR)."""
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("pymupdf required for PDF text extraction") from exc

    doc = fitz.open(pdf_path)
    parts: list[str] = []
    for page in doc:
        block = page.get_text("text").strip()
        if block:
            parts.append(block)
    doc.close()
    text = "\n\n".join(parts).strip()
    if len(text) < 50:
        raise ValueError(
            "PDF has little or no embedded text (likely a scan). "
            "Set CATTS_OCR_ENGINE=unlimited and a GPU worker for OCR."
        )
    return text


def docx_to_text(docx_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx required for DOCX support") from exc
    doc = Document(str(docx_path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    text = "\n\n".join(parts).strip()
    if len(text) < 10:
        raise ValueError("DOCX appears empty or unreadable")
    return text


def extract_text(source_path: Path) -> str:
    suffix = source_path.suffix.lower()
    if suffix in (".txt", ".md", ".markdown"):
        return source_path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        return docx_to_text(source_path)
    if suffix == ".epub":
        try:
            return epub_to_text(source_path)
        except Exception:
            return _epub_fallback_zip(source_path)
    if suffix == ".pdf":
        return pdf_to_text(source_path)
    raise ValueError(f"Unsupported input format: {suffix}")
