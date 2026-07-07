import io
import os
import shutil
import time


os.environ.setdefault("CATTS_DATA_DIR", "data/_tmp_test_api_uploads")

from fastapi.testclient import TestClient

from api.main import app
from config import DATA_DIR
from db import init_db


def setup_module():
    if DATA_DIR.exists():
        shutil.rmtree(DATA_DIR)
    init_db()


def _post_book(filename: str, payload: bytes, mime: str):
    client = TestClient(app)
    response = client.post(
        "/jobs/audiobook",
        files={"file": (filename, io.BytesIO(payload), mime)},
        data={
            "generate_audio": "false",
            "chapter_mode": "number",
            "title": "Upload Test",
            "lang": "en",
        },
    )
    assert response.status_code == 200, response.text
    job_id = response.json()["job_id"]
    deadline = time.time() + 10
    status = None
    while time.time() < deadline:
        status_response = client.get(f"/jobs/{job_id}")
        assert status_response.status_code == 200, status_response.text
        status = status_response.json()
        if status["status"] in {"done", "failed", "cancelled"}:
            break
        time.sleep(0.1)

    assert status is not None
    assert status["status"] == "done", status
    assert status["manuscript_ready"] is True

    files_response = client.get(f"/jobs/{job_id}/files")
    assert files_response.status_code == 200, files_response.text
    files = files_response.json()
    assert files["manuscript_md"]
    assert files["manuscript_txt"]


def test_txt_upload_produces_manuscript():
    _post_book(
        "book.txt",
        b"Chapter 1\nThis is a TXT upload test for CATTS.",
        "text/plain",
    )


def test_docx_upload_produces_manuscript():
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("This is a DOCX upload test for CATTS.")
    buffer = io.BytesIO()
    doc.save(buffer)
    _post_book(
        "book.docx",
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def test_pdf_upload_with_embedded_text_produces_manuscript():
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "Chapter 1\n"
        "This is a PDF upload test for CATTS with enough embedded text "
        "to pass the scanned-PDF guard and produce a manuscript.",
    )
    payload = doc.tobytes()
    doc.close()
    _post_book("book.pdf", payload, "application/pdf")


def test_epub_upload_produces_manuscript():
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("catts-test")
    book.set_title("CATTS EPUB Test")
    book.set_language("en")
    chapter = epub.EpubHtml(title="Chapter 1", file_name="chap_01.xhtml", lang="en")
    chapter.content = "<h1>Chapter 1</h1><p>This is an EPUB upload test for CATTS.</p>"
    book.add_item(chapter)
    book.toc = (epub.Link("chap_01.xhtml", "Chapter 1", "chap1"),)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    _post_book("book.epub", buffer.getvalue(), "application/epub+zip")
